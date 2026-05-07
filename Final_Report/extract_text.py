from docx import Document
import json

doc = Document(r"Final_Report/Final Year Project Report Template.docx")

output = []
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        output.append({
            "index": i,
            "style": para.style.name if para.style else "Normal",
            "text": para.text.strip()
        })

# Also extract tables
tables_output = []
for t_idx, table in enumerate(doc.tables):
    table_data = []
    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        table_data.append(row_data)
    tables_output.append({"table_index": t_idx, "rows": table_data})

with open("Final_Report/extracted_content.json", "w", encoding="utf-8") as f:
    json.dump({"paragraphs": output, "tables": tables_output}, f, indent=2, ensure_ascii=False)

print(f"Extracted {len(output)} paragraphs and {len(tables_output)} tables")
